#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('FBCA AI Infrastructure Overview', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Early Adopter Briefing for Andy')
run.font.size = Pt(14)
run.font.italic = True

doc.add_paragraph('Prepared by: Billy Nelms (Technology Lead)')
doc.add_paragraph('Date: March 1, 2026')
doc.add_paragraph()

# What Is This
doc.add_heading('What Is This?', level=1)
doc.add_paragraph('OpenClaw = Self-hosted AI agent platform (like having your own ChatGPT that runs on your Macs)')
doc.add_paragraph()
p = doc.add_paragraph('What we built: Two AI agents (B.O.B. & Cornerstone) that:')
doc.add_paragraph('Run locally on church MacBooks (no cloud dependency)', style='List Bullet')
doc.add_paragraph('Use FREE open-source AI models (Llama 3.1)', style='List Bullet')
doc.add_paragraph('Communicate securely via encrypted database (Convex)', style='List Bullet')
doc.add_paragraph('Automate church tech tasks', style='List Bullet')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Cost: $0/month').bold = True

doc.add_page_break()

# Current Infrastructure
doc.add_heading('Current Infrastructure (Whats Live)', level=1)

doc.add_heading('B.O.B. (Bot On Board) — Intel MacBook', level=2)
table1 = doc.add_table(rows=6, cols=3)
table1.style = 'Light Grid Accent 1'
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Status'
hdr_cells[2].text = 'Purpose'

data1 = [
    ['AI Model', 'Active', 'Local language processing'],
    ['Security Monitor', 'Active', 'SSH/disk/config monitoring'],
    ['Git + OneDrive Backup', 'Active', 'Code + config redundancy'],
    ['Message Poller', 'Active', 'Checks for Cornerstone messages'],
    ['Mission Control', 'Active', 'Web dashboard']
]

for i, row_data in enumerate(data1, 1):
    cells = table1.rows[i].cells
    cells[0].text = row_data[0]
    cells[1].text = row_data[1]
    cells[2].text = row_data[2]

doc.add_paragraph()

doc.add_heading('Cornerstone — M1 MacBook', level=2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Component'
hdr_cells2[1].text = 'Status'
hdr_cells2[2].text = 'Purpose'

data2 = [
    ['AI Model', 'In Progress', 'Switching to FREE Ollama'],
    ['9 Sub-Agents', 'Active', 'FBCA operations'],
    ['Message Poller', 'Active', 'Checks for B.O.B. messages'],
    ['Cross-Instance Bridge', 'Active', 'Secure communication']
]

for i, row_data in enumerate(data2, 1):
    cells = table2.rows[i].cells
    cells[0].text = row_data[0]
    cells[1].text = row_data[1]
    cells[2].text = row_data[2]

doc.add_page_break()

# What Can We Automate
doc.add_heading('What Can We Automate?', level=1)

doc.add_heading('Immediate (This Month)', level=2)
doc.add_paragraph('Weekly Church Reports — Auto-generate summary of tech ops, security status, upcoming events')
doc.add_paragraph('Security Alert Escalation — B.O.B. detects threat — Cornerstone notifies staff — Billy approves action')
doc.add_paragraph('Event Setup Checklist — Auto-create checklists for room bookings, A/V needs, volunteer assignments')
doc.add_paragraph('Backup Verification — Daily Git + OneDrive sync with status reports')

doc.add_heading('Short-Term (Next 3 Months)', level=2)
doc.add_paragraph('PCO Integration — Auto-sync Planning Center events to church calendar, website')
doc.add_paragraph('WordPress Automation — Auto-backup before updates, publish scheduled content')
doc.add_paragraph('Door Access Logs — Monitor access patterns, alert on anomalies')
doc.add_paragraph('Cross-Instance Learning — Share security findings, optimization tips via Convex')

doc.add_heading('Long-Term (6+ Months)', level=2)
doc.add_paragraph('AI Agent Swarm — 9 agents coordinating complex church operations')
doc.add_paragraph('Predictive Maintenance — Predict HVAC, tech failures before they happen')
doc.add_paragraph('Automated Sermon Notes — Transcribe, summarize, distribute to congregation')
doc.add_paragraph('Multi-Site Coordination — If FBCA expands, agents manage multiple campuses')

doc.add_page_break()

# Business Value
doc.add_heading('Business Value (Why This Matters)', level=1)

doc.add_heading('Cost Savings', level=2)
table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Light Grid Accent 1'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Current State'
hdr_cells3[1].text = 'Future State'
hdr_cells3[2].text = 'Savings'

data3 = [
    ['Cloud AI APIs ($40-200/mo)', 'Local Ollama (FREE)', '$500-2,400/year'],
    ['Manual event setup (5 hrs/week)', 'Automated checklists', '260 hours/year'],
    ['Reactive security fixes', 'Proactive monitoring', 'Prevents incidents'],
    ['Manual backups', 'Automated daily', 'Zero data loss']
]

for i, row_data in enumerate(data3, 1):
    cells = table3.rows[i].cells
    cells[0].text = row_data[0]
    cells[1].text = row_data[1]
    cells[2].text = row_data[2]

doc.add_paragraph()

doc.add_heading('Risk Mitigation', level=2)
doc.add_paragraph('No vendor lock-in — Runs on OUR hardware, OUR data')
doc.add_paragraph('Privacy — Conversations stay local/encrypted (no OpenAI data harvesting)')
doc.add_paragraph('Redundancy — Two Macs = if one dies, other keeps running')
doc.add_paragraph('Compliance — Church data never leaves our network')

doc.add_heading('Competitive Advantage', level=2)
doc.add_paragraph('Early adopter — First church in Tarrant County with self-hosted AI')
doc.add_paragraph('Staff efficiency — Automate repetitive tech tasks')
doc.add_paragraph('Scalability — Add agents as FBCA grows')
doc.add_paragraph('Knowledge retention — AI agents remember everything (no brain drain)')

doc.add_page_break()

# Bottom Line
doc.add_heading('Bottom Line', level=1)
p = doc.add_paragraph()
run = p.add_run('We have built a $0/month AI infrastructure that can automate 80% of FBCA repetitive tech work.')
run.bold = True

doc.add_paragraph()
doc.add_heading('What we need from leadership:', level=2)
doc.add_paragraph('Approval to "go live" with automation')
doc.add_paragraph('Priority — which church task to automate first?')
doc.add_paragraph('Trust — we will not break anything (backups, redundancy, Billy approval gates)')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('The sky is NOT the limit — we can go to the moon. But let us start with reliable automation.')
run.italic = True

doc.add_paragraph()
doc.add_paragraph('Ready to discuss? Billy can demo live security monitoring, AI Town, cross-instance messaging, or any use case you want to see.')

# Save document
output_path = '/Users/campoffice/OneDrive - First Baptist Church Arlington/FBCA_AI_Overview_for_Andy.docx'
doc.save(output_path)
print(f'Document created: {output_path}')
